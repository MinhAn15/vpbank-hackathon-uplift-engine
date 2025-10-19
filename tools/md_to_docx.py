"""Simple Markdown to DOCX converter for the UpliftEngine_Presentation_Summary.md

This script performs lightweight conversions: headings to Word headings, paragraphs,
and simple lists. It's intentionally small and deterministic for the hackathon.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from pathlib import Path


def add_heading(doc, text, level):
    # Map markdown heading level to Word heading and style it
    h = doc.add_heading(text, level=level)
    # Apply Times New Roman for headings
    for run in h.runs:
        run.font.name = 'Times New Roman'
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_list(doc, lines):
    for line in lines:
        clean = line.lstrip('-').strip()
        p = doc.add_paragraph(clean, style='List Bullet')
        p.runs[0].font.size = Pt(11)


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()
    # Set default Normal style to Arial 11 for body text
    try:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
    except Exception:
        pass
    # Cover page
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run('Uplift Engine 2.1 — Project Summary\n')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(26)
    run.bold = True
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_run = meta.add_run('Author: Team Uplift\nDate: 2025-10-19')
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(11)
    doc.add_page_break()

    # Table of Contents (Word will render it when the doc is opened and TOC updated)
    toc_para = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    # Use double backslashes so Python does not interpret \u as a unicode escape
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    toc_para._p.append(fld)
    doc.add_page_break()

    lines = md_path.read_text(encoding='utf-8').splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        # Headings
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            add_heading(doc, text, level=min(level, 3))
            i += 1
            continue
        # Ordered or unordered lists
        if re.match(r'^\s*[-\*]\s+', line):
            buffer = []
            while i < len(lines) and re.match(r'^\s*[-\*]\s+', lines[i]):
                buffer.append(lines[i])
                i += 1
            add_list(doc, buffer)
            continue
        # Markdown table
        if re.match(r'^\|', line):
            # collect contiguous table lines
            tbl_lines = []
            while i < len(lines) and re.match(r'^\|', lines[i]):
                tbl_lines.append(lines[i])
                i += 1
            render_table(doc, tbl_lines)
            continue
        # Paragraph
        # If we hit the Slide-ready section header, render the following block as a slide
        if re.match(r'^Slide-ready one-page summary', line):
            i += 1
            render_slide_summary(doc, lines, i)
            # advance index past the slide block: find next blank line after summary end
            while i < len(lines) and lines[i].strip() != '':
                i += 1
            continue
        add_paragraph(doc, line)
        i += 1

    doc.save(docx_path)


def render_table(doc, tbl_lines):
    # Simple pipe-separated table renderer. First line = header, second = separator
    rows = [ [cell.strip() for cell in re.split(r'\|', ln)[1:-1]] for ln in tbl_lines]
    if len(rows) < 2:
        return
    header = rows[0]
    data = rows[2:] if re.match(r'^[\s\|:-]+$', rows[1][0]) else rows[1:]
    table = doc.add_table(rows=1+len(data), cols=len(header))
    table.style = 'Light List'
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(header):
        hdr_cells[j].text = h
    for i, row in enumerate(data, start=1):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell


def render_slide_summary(doc, lines, start_idx):
    # Render a compact, centered slide-style page for executives
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('Executive Summary — Pilot Ask & Impact\n')
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()
    # collect up to 12 non-empty lines or until next section
    i = start_idx
    bullets = []
    while i < len(lines):
        ln = lines[i].strip()
        if not ln:
            break
        bullets.append(ln)
        i += 1
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')
        p.runs[0].font.size = Pt(12)
    doc.add_page_break()


if __name__ == '__main__':
    import sys
    root = Path(__file__).resolve().parents[1]
    md = root / 'docs' / 'UpliftEngine_Presentation_Summary.md'
    out = root / 'docs' / 'UpliftEngine_Presentation_Summary.docx'
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out = Path(sys.argv[2])
    print(f'Converting {md} -> {out}')
    md_to_docx(md, out)
    print('Done')
